import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import Config
import re

class DocumentManager:
    """
    Manages document generation and updates from chat conversations.
    """
    
    def __init__(self):
        self.documents_folder = Config.DOCUMENTS_FOLDER
        os.makedirs(self.documents_folder, exist_ok=True)
    
    def update_document(self, session_id, query, answer, sources, is_deep_dive=False):
        doc_data_path = os.path.join(self.documents_folder, f"{session_id}.json")
        if os.path.exists(doc_data_path):
            with open(doc_data_path, 'r', encoding='utf-8') as f:
                doc_data = json.load(f)
        else:
            doc_data = {
                'session_id': session_id,
                'created_at': datetime.now().isoformat(),
                'topics': []
            }
        topic = self._extract_topic(query)
        topic_entry = None
        for t in doc_data['topics']:
            if self._is_similar_topic(t['topic'], topic):
                topic_entry = t
                break
        if topic_entry:
            if is_deep_dive:
                topic_entry['content'] = answer
                topic_entry['sources'] = sources
                topic_entry['updated_at'] = datetime.now().isoformat()
        else:
            doc_data['topics'].append({
                'topic': topic,
                'query': query,
                'content': answer,
                'sources': sources,
                'created_at': datetime.now().isoformat(),
                'updated_at': datetime.now().isoformat()
            })
        with open(doc_data_path, 'w', encoding='utf-8') as f:
            json.dump(doc_data, f, indent=2, ensure_ascii=False)

    def generate_preview(self, session_id):
        doc_data_path = os.path.join(self.documents_folder, f"{session_id}.json")
        if not os.path.exists(doc_data_path):
            return "No document available yet. Start a conversation to generate content."
        with open(doc_data_path, 'r', encoding='utf-8') as f:
            doc_data = json.load(f)
        html = f"<h1>Company Research Report</h1>\n"
        html += f"<p><em>Generated: {datetime.fromisoformat(doc_data['created_at']).strftime('%B %d, %Y at %I:%M %p')}</em></p>\n"
        html += "<hr/>\n"
        for idx, topic in enumerate(doc_data['topics'], 1):
            html += f"<h2>{idx}. {topic['topic']}</h2>\n"
            html += f"<p><strong>Query:</strong> {topic['query']}</p>\n"
            clean_answer = re.sub(r'\[Source \d+\]', '', topic['content'])
            html_answer = format_for_html_preview(clean_answer)
            html += html_answer
            # Add sources
            if topic['sources']:
                html += "<h3>Sources</h3>\n<ul class='source-list'>"
                for source in topic['sources'][:5]:
                    html += f"<li><a href='{source['url']}' target='_blank'>{source['title']}</a></li>"
                html += "</ul>"
            html += "<hr/>\n"
        return html

    def generate_docx(self, session_id):
        doc_data_path = os.path.join(self.documents_folder, f"{session_id}.json")
        if not os.path.exists(doc_data_path):
            return None
        with open(doc_data_path, 'r', encoding='utf-8') as f:
            doc_data = json.load(f)
        doc = Document()
        title = doc.add_heading('Company Research Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Generated: {datetime.fromisoformat(doc_data['created_at']).strftime('%B %d, %Y at %I:%M %p')}").italic = True
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        for idx, topic in enumerate(doc_data['topics'], 1):
            doc.add_heading(f"{idx}. {topic['topic']}", 1)
            query_para = doc.add_paragraph()
            query_para.add_run("Query: ").bold = True
            query_para.add_run(topic['query'])
            doc.add_paragraph()
            clean_content = re.sub(r'\[Source \d+\]', '', topic['content'])
            format_for_docx(doc, clean_content)  # Enhanced formatting!
            if topic['sources']:
                doc.add_heading('Sources', 2)
                for source in topic['sources'][:5]:
                    source_para = doc.add_paragraph(style='List Bullet')
                    source_para.add_run(source['title']).bold = True
                    source_para.add_run(f"\n{source['url']}")
            doc.add_page_break()
        docx_path = os.path.join(self.documents_folder, f"{session_id}.docx")
        doc.save(docx_path)
        return docx_path

    def _extract_topic(self, query):
        words = query.split()[:8]
        topic = ' '.join(words)
        if len(query.split()) > 8:
            topic += '...'
        return topic

    def _is_similar_topic(self, topic1, topic2):
        words1 = set(topic1.lower().split())
        words2 = set(topic2.lower().split())
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        if not union:
            return False
        similarity = len(intersection) / len(union)
        return similarity > 0.5

# Helper: Format for HTML preview
def format_for_html_preview(answer):
    answer = re.sub(r'\*\*([^\*]+)\*\*', r'<h2>\1</h2>', answer)
    answer = re.sub(r'(?:^|\n)\s*[•\-]\s*(.+?)(?=\n|$)', r'<li>\1</li>', answer)
    answer = re.sub(r'((?:<li>.*?</li>\s*)+)', r'<ul>\1</ul>', answer, flags=re.DOTALL)
    parts = re.split(r'(\n+)', answer)
    html = ''
    for part in parts:
        part = part.strip()
        if part.startswith('<h2>') or part.startswith('<ul>') or not part:
            html += part
        else:
            html += f'<p>{part}</p>'
    html = re.sub(r'<p>\s*</p>', '', html)
    return html

# Helper: Format for DOCX export
def format_for_docx(doc, content):
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if re.match(r'\*\*.+\*\*', line):
            heading = re.sub(r'\*\*', '', line)
            doc.add_heading(heading, level=2)
        elif line.startswith('•') or line.startswith('-'):
            doc.add_paragraph(line[1:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line)
